import streamlit as st
import google.generativeai as genai
from PIL import Image
import io
import base64
from docx import Document

# Page configuration
st.set_page_config(
    page_title="GSM-R Network Analyzer",
    page_icon="🚄",
    layout="wide"
)

# Configuration 
GEMINI_API_KEY = st.secrets["GEMINI_API_KEY"]

# --- Custom CSS ---
st.markdown("""
<style>
    .main-header {
        text-align: center;
        color: #f4730e;
        font-size: 2.5rem;
        font-weight: bold;
        margin-bottom: 2rem;
        padding: 1rem;
        background: linear-gradient(90deg, #fff1e6, #ffffff);
        border-radius: 10px;
        box-shadow: 0 2px 4px rgba(244, 115, 14, 0.1);
        border: 2px solid #f4730e20;
    }
    .input-section {
        background: #fff8f2;
        padding: 2rem;
        border-radius: 10px;
        margin: 1rem 0;
        border: 1px solid #f4730e30;
    }
    .result-section {
        background: #fff1e6;
        padding: 2rem;
        border-radius: 10px;
        margin: 1rem 0;
        border-left: 5px solid #f4730e;
    }
    .stButton > button {
        background-color: #f4730e !important;
        color: white !important;
        font-weight: bold !important;
        border-radius: 8px !important;
        border: none !important;
        padding: 0.75rem 2rem !important;
        font-size: 1.1rem !important;
        transition: all 0.3s ease !important;
    }
    .stButton > button:hover {
        background-color: #d85a0a !important;
        transform: translateY(-2px) !important;
        box-shadow: 0 4px 8px rgba(244, 115, 14, 0.3) !important;
    }
</style>
""", unsafe_allow_html=True)

# --- Utility functions ---

def display_logo():
    try:
        logo = Image.open("Logo_ONCF.svg.png")
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            st.image(logo, width=300)
            st.markdown(
                '<p style="text-align: center; color: #666; margin-top: 0.5rem; font-size: 0.9rem;">Office National des Chemins de Fer</p>',
                unsafe_allow_html=True
            )
    except:
        st.markdown("""
        <div style="text-align: center; margin: 2rem 0;">
            <div style="background: linear-gradient(45deg, #f4730e, #ff8c42); 
                        color: white; 
                        font-size: 2rem; 
                        font-weight: bold; 
                        padding: 1rem 2rem; 
                        border-radius: 10px; 
                        display: inline-block;
                        box-shadow: 0 4px 8px rgba(244, 115, 14, 0.3);">
                🚄 ONCF
            </div>
            <p style="color: #666; margin-top: 0.5rem; font-size: 0.9rem;">
                Office National des Chemins de Fer
            </p>
        </div>
        """, unsafe_allow_html=True)

def get_gemini_response(prompt, api_key):
    try:
        genai.configure(api_key=api_key)
        model_names = [
            'gemini-2.0-flash',
            'gemini-1.5-flash',
            'gemini-1.5-pro', 
            'gemini-1.0-pro'
        ]
        model = None
        for model_name in model_names:
            try:
                model = genai.GenerativeModel(model_name)
                break
            except:
                continue
        if model is None:
            return "Could not find suitable model. Check your API key or available models."

        system_prompt = """
Vous êtes un expert hautement qualifié en réseau GSM-R (Global System for Mobile Communications - Railway), spécialisé dans les communications ferroviaires à grande vitesse.

**Règle principale :**
- Si la question ou la demande n’est pas liée au GSM-R ou ne contient aucun élément technique du domaine GSM-R (paramètres radio, handover, BSC, Cell ID, fréquence, puissance, etc.), répondez uniquement :
  "Veuillez saisir les détails de la déconnexion."
  et rien d’autre.

**Sinon :**
- Donnez des recommandations *courtes*, *claires* et *sous forme de points-bullets* uniquement sur 1 ou 2 paramètres maximum à optimiser à la fois (pas toute la liste).

Pour chaque paramètre recommandé :
- Indiquez :
  - Nom du paramètre
  - Valeur actuelle vs valeur recommandée
  - Justification technique (1 seule phrase)
  - Niveau de priorité (Critique / Haute / Moyenne / Basse)
  - Impact attendu (1 seule phrase)

**Format de sortie :**
- Utilisez des puces claires (•)
- Pas de paragraphes longs, pas d'introduction ni de conclusion
- Maximum 2 paramètres à la fois

**Contexte :**
- Ligne LGV TGV entre Rabat et Tanger au Maroc.
- Cellules principales 201-233 (BSC Rabat) ; cellules de secours 301-333 (BSC Kénitra).

**Exemples :**

• Penalty Stop Level Threshold (Cell ID 202 et Neighbor 2G Cell ID 201): 20 -> 25
  - Justification : Augmente la robustesse HO en cas de perte de signal rapide.
  - Priorité : Haute
  - Impact attendu : Moins de coupures sur zones à forte vitesse.

• PBGT HO Threshold (Source Cell ID 202 -> Neighbor 2G Cell ID 201): 82 -> 85
  - Justification : Améliore la précision du déclenchement HO.
  - Priorité : Moyenne
  - Impact attendu : Réduction des HO prématurés.

• Penalty Timer Length (Source Cell ID 202 -> Neighbor 2G Cell ID 201): 10 -> 15
  - Justification : Allonge la fenêtre d'évaluation HO pour stabiliser les décisions.
  - Priorité : Moyenne
  - Impact attendu : Moins de ping-pong HO.

• PBGT HO Threshold (Source Cell ID 202 -> Neighbor 2G Cell ID 203): 64 -> 60
  - Justification : Rend la transition HO vers la cellule voisine plus réactive.
  - Priorité : Haute
  - Impact attendu : Moins de coupures lors des changements de BSC.

• BQ HO Margin (Source Cell ID 202 -> Neighbor 2G Cell ID 203): 69 -> 72
  - Justification : Ajuste la marge HO pour zones inter-BSC.
  - Priorité : Moyenne
  - Impact attendu : Meilleure continuité sur la frontière Rabat/Kénitra.

• Fréquence couche 201: 955 -> 956
  - Justification : Réduction des interférences adjacentes.
  - Priorité : Critique
  - Impact attendu : Amélioration de la qualité radio sur la section Rabat.

• Fréquence couche 301: 959 -> 974
  - Justification : Séparation de canal pour éviter les chevauchements.
  - Priorité : Critique
  - Impact attendu : Moins d'interférences avec la voie principale.

• Réduction de la puissance de 2dB pour la couche 201 : 40dBm -> 38dBm
  - Justification : Limite les interférences inter-cellules sur la ligne.
  - Priorité : Haute
  - Impact attendu : Meilleur équilibrage de couverture.

• Réduction de la puissance de 2dB pour la couche 301 : 40dBm -> 38dBm
  - Justification : Réduction des interférences avec les cellules voisines.
  - Priorité : Moyenne
  - Impact attendu : Couverture plus propre sur la ligne de secours.
"""



        full_prompt = f"{system_prompt}\n\nUser Query: {prompt}"
        response = model.generate_content(full_prompt)
        return response.text
    except Exception as e:
        return f"Error generating response: {str(e)}"

def export_to_word(text):
    doc = Document()
    doc.add_heading('GSM-R Network Analysis & Recommendations', 0)
    doc.add_paragraph(text)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Main App ---
def main():
    # Initialize session state
    if 'result_text' not in st.session_state:
        st.session_state.result_text = ""

    # Header
    display_logo()
    st.markdown('<h1 class="main-header">GSM-R Network Disconnection Analysis System</h1>', unsafe_allow_html=True)
    st.markdown('<p style="text-align: center; color: #666; font-size: 1.1rem; margin-bottom: 2rem;">LGV TGV Rabat-Tangier High-Speed Line Network Optimization Tool</p>', unsafe_allow_html=True)

    # Sidebar
    with st.sidebar:
        st.markdown('<div class="header-accent">📊 Network Information</div>', unsafe_allow_html=True)
        st.info("""
        **Main Route (BSC Rabat)**
        - Cells: 201-233
        **Backup Route (BSC Kenitra)**
        - Cells: 301-333
        **Line**: LGV TGV Rabat-Tangier
        """)
        st.markdown('<div class="header-accent">🔧 System Status</div>', unsafe_allow_html=True)
        if GEMINI_API_KEY != "YOUR_API_KEY_HERE":
            st.success("✅ API Configuration: Active")
        else:
            st.error("❌ API Configuration: Not Set")
        st.info("🌐 GSM-R Expert System: Online")

    # Input Section
    st.markdown('<div class="input-section">', unsafe_allow_html=True)
    st.subheader("🔍 Disconnection Event Analysis")
    user_input = st.text_area(
        "Describe the disconnection event details:",
        placeholder="Please provide detailed information about the disconnection event...",
        height=200
    )
    st.markdown('</div>', unsafe_allow_html=True)

    # Buttons
    col_gen, col_exp = st.columns(2)
    with col_gen:
        if st.button("🔬 Generate Network Analysis & Recommendations", use_container_width=True):
            if GEMINI_API_KEY == "YOUR_API_KEY_HERE":
                st.error("⚠️ Please configure your Gemini API key in the code before using the application.")
            elif user_input.strip() == "":
                st.warning("⚠️ Please enter details about the disconnection event.")
            else:
                with st.spinner("Generating recommendations..."):
                    result = get_gemini_response(user_input, GEMINI_API_KEY)
                    st.session_state.result_text = result

    with col_exp:
        if st.button("📝 Export Report to Word", use_container_width=True):
            if st.session_state.result_text.strip() == "":
                st.warning("⚠️ No analysis to export yet. Please generate it first.")
            else:
                buffer = export_to_word(st.session_state.result_text)
                st.download_button(
                    label="📥 Download Word Report",
                    data=buffer,
                    file_name="GSMR_Network_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

    # Results Section
    if st.session_state.result_text.strip():
        st.markdown('<div class="result-section">', unsafe_allow_html=True)
        st.subheader("📑 Analysis & Recommendations")
        st.markdown(st.session_state.result_text)
        st.markdown('</div>', unsafe_allow_html=True)

# --- Run app ---
if __name__ == "__main__":
    main()
